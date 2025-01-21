from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from docx import Document
from flask_sqlalchemy import SQLAlchemy
import os
import logging
import pandas as pd
from datetime import datetime

# Initialize Flask app
app = Flask(__name__)
app.secret_key = "supersecretkey"
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///database.db"
app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# Initialize Database
db = SQLAlchemy(app)

# Logging setup
logging.basicConfig(filename="app.log", level=logging.ERROR, 
                    format="%(asctime)s - %(levelname)s - %(message)s")

# Database Model
class Applicant(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    age = db.Column(db.Integer, nullable=False)
    address = db.Column(db.String(200), nullable=False)
    phone = db.Column(db.String(15), nullable=False)

# Create Database
with app.app_context():
    db.create_all()

# Helper function to detect placeholders in Word template
def get_placeholders(template_path):
    doc = Document(template_path)
    placeholders = set()
    for paragraph in doc.paragraphs:
        placeholders.update([word.strip("{}") for word in paragraph.text.split() if word.startswith("{") and word.endswith("}")])
    return list(placeholders)

# Fill Word template
def fill_form(template_path, output_path, data):
    try:
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
        doc.save(output_path)
    except Exception as e:
        logging.error(f"Error filling form: {e}")
        raise

# Routes
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/submit", methods=["POST"])
def submit_form():
    try:
        # Get form data
        name = request.form["name"]
        age = request.form["age"]
        address = request.form["address"]
        phone = request.form["phone"]

        # Save data to database
        applicant = Applicant(name=name, age=age, address=address, phone=phone)
        db.session.add(applicant)
        db.session.commit()

        # Create folder with today's date
        today = datetime.now().strftime("%Y-%m-%d")
        output_folder = os.path.join("filled_forms", today)
        os.makedirs(output_folder, exist_ok=True)

        # Fill the form
        template_path = "template.docx"
        output_path = os.path.join(output_folder, f"{name}_form.docx")
        fill_form(template_path, output_path, {"Name": name, "Age": age, "Address": address, "Phone": phone})

        return send_file(output_path, as_attachment=True)

    except Exception as e:
        logging.error(f"Error processing form: {e}")
        flash("An error occurred. Check logs for details.")
        return redirect(url_for("index"))

@app.route("/batch", methods=["GET", "POST"])
def batch_form():
    if request.method == "POST":
        try:
            # Upload and process CSV
            file = request.files["file"]
            if not file:
                flash("Please upload a CSV file.")
                return redirect(url_for("batch_form"))

            file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(file_path)
            data = pd.read_csv(file_path)

            # Create folder with today's date
            today = datetime.now().strftime("%Y-%m-%d")
            output_folder = os.path.join("filled_forms", today)
            os.makedirs(output_folder, exist_ok=True)

            # Validate and fill forms
            template_path = "template.docx"

            for index, row in data.iterrows():
                row_data = row.to_dict()

                # Clean and validate row data
                if "Name" not in row_data or not row_data["Name"]:
                    logging.error(f"Missing 'Name' in row {index + 1}")
                    continue

                name = row_data["Name"].strip().replace(" ", "_")
                output_path = os.path.join(output_folder, f"{name}_form.docx")

                # Fill the form
                try:
                    fill_form(template_path, output_path, row_data)
                except Exception as e:
                    logging.error(f"Error filling form for {name}: {e}")
                    continue

            flash(f"Batch forms generated successfully in folder: {today}")
            return redirect(url_for("index"))

        except Exception as e:
            logging.error(f"Error processing batch forms: {e}")
            flash("An error occurred. Check logs for details.")
            return redirect(url_for("batch_form"))

    return render_template("batch.html")

# Route to show database details
@app.route("/history")
def history():
    try:
        # Fetch all applicants from the database, ordered by ID in descending order
        applicants = Applicant.query.order_by(Applicant.id.desc()).all()
        return render_template("history.html", applicants=applicants)
    except Exception as e:
        logging.error(f"Error fetching history: {e}")
        flash("An error occurred while fetching the history.")
        return redirect(url_for("index"))



if __name__ == "__main__":
    app.run(debug=True)
